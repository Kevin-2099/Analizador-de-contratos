import gradio as gr
import re
import html
import csv
import io
import tempfile
from datetime import datetime
from difflib import SequenceMatcher
from langdetect import detect

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt

# ── Optional dependencies (graceful fallback) ─────────────────────────────────
try:
    import pdfplumber
    PDF_READ = True
except ImportError:
    PDF_READ = False

try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import PatternFill, Font, Alignment
    from openpyxl.utils import get_column_letter
    XLSX_SUPPORT = True
except ImportError:
    XLSX_SUPPORT = False


# ── Supported languages ───────────────────────────────────────────────────────
LANG_LABELS = {"auto": "🌐 Auto", "es": "🇪🇸 Español", "en": "🇬🇧 English",
               "fr": "🇫🇷 Français", "pt": "🇵🇹 Português"}

# ── Keywords per clause per language ─────────────────────────────────────────
KEYWORDS: dict[str, dict[str, list[str]]] = {
    "es": {
        "pagos":          ["pago","pagará","abonará","monto","importe","tarifa","honorario",
                           "remuneración","cuota","factura","transferencia","depósito","salario",
                           "bono","comisión","costo"],
        "penalizaciones": ["penalización","multa","interés moratorio","sanción","recargo",
                           "compensación","indemnización","daños","perjuicio"],
        "obligaciones":   ["deberá","se obliga","obligación","compromete","cumplir","garantizar",
                           "proveer","asegurar","entregar","informar"],
        "confidencialidad":["confidencialidad","no divulgación","NDA","información confidencial",
                            "datos sensibles","privada","secreto"],
        "terminación":    ["terminación","cancelación","rescisión","extinción","finalización",
                           "dar por terminado"],
        "jurisdiccion":   ["jurisdicción","ley aplicable","tribunales competentes","fuero",
                           "arbitraje","arbitraje vinculante","legislación aplicable",
                           "sometidas a los tribunales"],
        "fuerza_mayor":   ["fuerza mayor","caso fortuito","causas ajenas a su voluntad",
                           "evento imprevisible","circunstancias extraordinarias","pandemia",
                           "desastre natural"],
        "propiedad_intelectual": ["propiedad intelectual","derechos de autor","marca registrada",
                           "patente","titularidad","propiedad industrial","derechos de propiedad",
                           "know-how","licencia de uso"],
        "limitacion_responsabilidad": ["limitación de responsabilidad","responsabilidad limitada",
                           "no será responsable","exención de responsabilidad","límite máximo de responsabilidad",
                           "daños indirectos","lucro cesante"],
        "no_competencia": ["no competencia","no competir","exclusividad","cláusula de no competencia",
                           "abstenerse de competir","actividad competidora"],
        "cesion":         ["cesión","ceder el contrato","transferir a terceros","subcontratar",
                           "sin consentimiento previo","cesión de derechos y obligaciones"],
    },
    "en": {
        "pagos":          ["payment","shall pay","fee","amount","rate","invoice","transfer",
                           "deposit","salary","bonus","remuneration"],
        "penalizaciones": ["penalty","fine","interest","surcharge","damages","indemnity"],
        "obligaciones":   ["shall","must","is obligated","required","comply","provide","deliver"],
        "confidencialidad":["confidentiality","non-disclosure","NDA","sensitive information",
                            "private","secret"],
        "terminación":    ["termination","cancellation","rescission","expiration","end of contract"],
        "jurisdiccion":   ["jurisdiction","governing law","competent courts","venue",
                           "arbitration","binding arbitration","applicable law"],
        "fuerza_mayor":   ["force majeure","act of god","beyond its reasonable control",
                           "unforeseeable event","pandemic","natural disaster"],
        "propiedad_intelectual": ["intellectual property","copyright","trademark","patent",
                           "ownership","proprietary rights","know-how","license to use"],
        "limitacion_responsabilidad": ["limitation of liability","limited liability",
                           "shall not be liable","disclaimer of liability","liability cap",
                           "indirect damages","consequential damages"],
        "no_competencia": ["non-compete","non-competition","exclusivity","competing business",
                           "refrain from competing"],
        "cesion":         ["assignment","assign this agreement","transfer to third parties",
                           "subcontract","without prior consent"],
    },
    "fr": {
        "pagos":          ["paiement","rémunération","montant","facture","honoraires","salaire"],
        "penalizaciones": ["pénalité","amende","sanction","dommages","indemnité"],
        "obligaciones":   ["devra","est tenu","obligation","s'engage","fournir","livrer"],
        "confidencialidad":["confidentialité","non-divulgation","information confidentielle","secret"],
        "terminación":    ["résiliation","annulation","expiration","fin du contrat"],
        "jurisdiccion":   ["juridiction","droit applicable","tribunaux compétents","arbitrage"],
        "fuerza_mayor":   ["force majeure","cas fortuit","événement imprévisible","pandémie"],
        "propiedad_intelectual": ["propriété intellectuelle","droits d'auteur","marque déposée",
                           "brevet","savoir-faire"],
        "limitacion_responsabilidad": ["limitation de responsabilité","ne sera pas responsable",
                           "dommages indirects"],
        "no_competencia": ["non-concurrence","exclusivité","activité concurrente"],
        "cesion":         ["cession","céder le contrat","transférer à des tiers","sous-traiter"],
    },
    "pt": {
        "pagos":          ["pagamento","remuneração","valor","fatura","honorários","salário"],
        "penalizaciones": ["penalidade","multa","sanção","danos","indenização"],
        "obligaciones":   ["deverá","obriga-se","obrigação","comprometer","fornecer","entregar"],
        "confidencialidad":["confidencialidade","não divulgação","informação confidencial","segredo"],
        "terminación":    ["rescisão","cancelamento","extinção","término do contrato"],
        "jurisdiccion":   ["jurisdição","lei aplicável","foro competente","arbitragem"],
        "fuerza_mayor":   ["força maior","caso fortuito","evento imprevisível","pandemia"],
        "propiedad_intelectual": ["propriedade intelectual","direitos autorais","marca registrada",
                           "patente","know-how"],
        "limitacion_responsabilidad": ["limitação de responsabilidade","não será responsável",
                           "danos indiretos"],
        "no_competencia": ["não concorrência","exclusividade","atividade concorrente"],
        "cesion":         ["cessão","ceder o contrato","transferir a terceiros","subcontratar"],
    },
}

# ── Risk keywords per language ────────────────────────────────────────────────
RIESGOS: dict[str, dict[str, list[str]]] = {
    "es": {"Bajo":["recargo"], "Moderado":["penalización","sanción"],
           "Alto":["incumplimiento","daños"], "Crítico":["indemnización","perjuicio"]},
    "en": {"Bajo":["surcharge"], "Moderado":["penalty","fine"],
           "Alto":["breach","damages"], "Crítico":["indemnity","loss"]},
    "fr": {"Bajo":["supplément"], "Moderado":["pénalité","amende"],
           "Alto":["manquement","dommages"], "Crítico":["indemnité","préjudice"]},
    "pt": {"Bajo":["acréscimo"], "Moderado":["penalidade","multa"],
           "Alto":["inadimplemento","danos"], "Crítico":["indenização","prejuízo"]},
}

# ── Abusive clause patterns ───────────────────────────────────────────────────
ABUSIVAS: dict[str, list[str]] = {
    "es": ["a sola discreción","sin previo aviso","en cualquier momento y sin causa",
           "prórroga automática","renuncia irrevocable","sin responsabilidad alguna",
           "según estime conveniente","en tiempo razonable","a su entera discreción",
           "sin limitación alguna","sin necesidad de notificación","modificar en cualquier momento",
           "renuncia a cualquier reclamo","exención total de responsabilidad",
           "cesión sin consentimiento","responsabilidad ilimitada del cliente"],
    "en": ["at its sole discretion","without prior notice","at any time without cause",
           "automatic renewal","irrevocable waiver","without any liability",
           "as it deems appropriate","within reasonable time","without limitation",
           "without notice","modify at any time","waives any claim",
           "total disclaimer of liability","assignment without consent"],
    "fr": ["à sa seule discrétion","sans préavis","renouvellement automatique",
           "sans responsabilité","dans un délai raisonnable"],
    "pt": ["a seu exclusivo critério","sem aviso prévio","renovação automática",
           "sem qualquer responsabilidade","em prazo razoável"],
}

# ── Date / money regex ────────────────────────────────────────────────────────
DATE_PATTERNS = [
    r'\b\d{1,2}\s+de\s+\w+\s+de\s+\d{4}\b',
    r'\b\d{1,2}/\d{1,2}/\d{2,4}\b',
    r'\b\d{4}-\d{2}-\d{2}\b',
    r'\b\d{1,3}\s+días?\b',
    r'\b\d{1,2}\s+meses?\b',
    r'\b\d{1,2}\s+años?\b',
    r'\bwithin\s+\d+\s+(?:calendar\s+)?days?\b',
    r'\b\d+\s+business\s+days?\b',
    r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{1,2},?\s+\d{4}\b',
]
MONEY_PATTERNS = [
    r'(?:USD|EUR|MXN|COP|ARS|GBP|BRL|€|\$|£)\s*[\d,\.]+',
    r'[\d,\.]+\s*(?:USD|EUR|MXN|COP|ARS|GBP|BRL|pesos?|euros?|dólares?|dollars?)',
]
PARTY_PATTERNS = [
    # denominado/referred to as → (company_name, alias)
    r'([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑa-záéíóúñ\s\.]{3,55}?),?\s+(?:con\s+\w[\w\s\-]+?,\s+)?(?:en adelante|hereinafter)\s+(?:denominad[oa]|referred to as)\s+"?([^",\n]{3,35})"?',
    # entre X y Y → real company names
    r'(?:entre|between)\s+(?:la\s+empresa\s+|el\s+señor\s+|la\s+señora\s+)?([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s\.]{3,45}?)\s*,.*?(?:\sy\s|\sand\s)(?:la\s+empresa\s+|el\s+señor\s+|la\s+señora\s+)?([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s\.]{3,45}?)\s*(?:,|\.)',
]
_CLEAN_PREFIX = re.compile(
    r'^(?:entre\s+|between\s+|y\s+|and\s+|la\s+empresa\s+|el\s+señor\s+|la\s+señora\s+|the\s+company\s+)',
    re.IGNORECASE
)

# ── Vigencia / vencimiento keywords per language ──────────────────────────────
VIGENCIA_KEYWORDS: dict[str, list[str]] = {
    "es": ["vigencia","vencimiento","expira","expiración","fecha de terminación",
           "válido hasta","validez del presente contrato","plazo de vigencia"],
    "en": ["expiration","expiry","valid until","term of this agreement","end date",
           "effective until"],
    "fr": ["expiration","valable jusqu","durée du contrat","date de fin"],
    "pt": ["vigência","expira","validade do contrato","data de término"],
}

MESES: dict[str, dict[str, int]] = {
    "es": {"enero":1,"febrero":2,"marzo":3,"abril":4,"mayo":5,"junio":6,"julio":7,
           "agosto":8,"septiembre":9,"setiembre":9,"octubre":10,"noviembre":11,"diciembre":12},
    "en": {"january":1,"february":2,"march":3,"april":4,"may":5,"june":6,"july":7,
           "august":8,"september":9,"october":10,"november":11,"december":12},
    "fr": {"janvier":1,"février":2,"mars":3,"avril":4,"mai":5,"juin":6,"juillet":7,
           "août":8,"septembre":9,"octobre":10,"novembre":11,"décembre":12},
    "pt": {"janeiro":1,"fevereiro":2,"março":3,"abril":4,"maio":5,"junho":6,"julho":7,
           "agosto":8,"setembro":9,"outubro":10,"novembro":11,"dezembro":12},
}

DURATION_PATTERN = re.compile(r'\b(\d{1,3})\s+(días?|meses?|años?|days?|months?|years?)\b', re.IGNORECASE)

# ── Templates by contract type ────────────────────────────────────────────────
WHY_LEGAL: dict[str, str] = {
    "pagos": "Define montos, plazos y forma de pago — su ausencia genera disputas sobre cuánto y cuándo se debe pagar.",
    "penalizaciones": "Establece consecuencias del incumplimiento — sin esto no hay disuasivo claro ni forma de exigir compensación.",
    "obligaciones": "Detalla qué debe hacer cada parte — es el núcleo del contrato; sin ella, las expectativas quedan ambiguas.",
    "confidencialidad": "Protege información sensible compartida entre las partes — crítica en NDAs y colaboraciones.",
    "terminación": "Explica cómo y cuándo puede finalizar el contrato — evita quedar atado indefinidamente o sin salida clara.",
    "jurisdiccion": "Define qué tribunales/leyes aplican en caso de disputa — sin esto un litigio puede ser costoso o incierto.",
    "fuerza_mayor": "Cubre eventos imprevisibles (desastres, pandemias) que impiden cumplir — sin ella cualquier incumplimiento externo podría considerarse falta.",
    "propiedad_intelectual": "Aclara quién es dueño de creaciones o desarrollos derivados — vital en servicios creativos o de desarrollo.",
    "limitacion_responsabilidad": "Limita el monto que una parte puede reclamar a la otra por daños — su ausencia expone a responsabilidad ilimitada.",
    "no_competencia": "Restringe que una parte compita con la otra — común en contratos laborales y de venta de negocios.",
    "cesion": "Regula si el contrato puede transferirse a un tercero sin consentimiento — clave en fusiones o subcontratación.",
}

TEMPLATES: dict[str, list[str]] = {
    "Genérico":                  ["pagos","penalizaciones","obligaciones","confidencialidad","terminación","jurisdiccion"],
    "NDA / Confidencialidad":    ["confidencialidad","obligaciones","terminación","jurisdiccion","propiedad_intelectual"],
    "Laboral":                   ["pagos","obligaciones","terminación","no_competencia","confidencialidad","jurisdiccion"],
    "Arrendamiento":             ["pagos","penalizaciones","obligaciones","terminación","jurisdiccion","fuerza_mayor"],
    "Prestación de servicios":   ["pagos","obligaciones","penalizaciones","terminación","propiedad_intelectual",
                                   "limitacion_responsabilidad","jurisdiccion"],
    "Compraventa":               ["pagos","obligaciones","penalizaciones","cesion","limitacion_responsabilidad",
                                   "jurisdiccion","fuerza_mayor"],
}

CLAUSE_EMOJIS = {"pagos":"💰","penalizaciones":"⚠️","obligaciones":"📌","confidencialidad":"🔒",
                  "terminación":"❌","jurisdiccion":"⚖️","fuerza_mayor":"🌪️",
                  "propiedad_intelectual":"©️","limitacion_responsabilidad":"🛡️",
                  "no_competencia":"🚧","cesion":"🔁"}
CLAUSE_COLORS = {"pagos":"#3b82f6","penalizaciones":"#f59e0b","obligaciones":"#8b5cf6",
                  "confidencialidad":"#10b981","terminación":"#ef4444","jurisdiccion":"#0ea5e9",
                  "fuerza_mayor":"#a855f7","propiedad_intelectual":"#ec4899",
                  "limitacion_responsabilidad":"#14b8a6","no_competencia":"#f97316","cesion":"#6366f1"}
RISK_COLORS   = {"Bajo":"#22c55e","Moderado":"#f59e0b","Alto":"#ef4444","Crítico":"#7c3aed"}
RISK_ICONS    = {"Bajo":"🟢","Moderado":"🟡","Alto":"🔴","Crítico":"💀"}
PESOS_DEFAULT = {"Bajo":1,"Moderado":2,"Alto":3,"Crítico":4}

# ═══════════════════════════════════════════════════════════════════════════════
# CORE ANALYSIS FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def generar_regex(palabras: list[str], ventana: int = 150) -> re.Pattern:
    escaped = [re.escape(w) for w in palabras]
    return re.compile(r".{0," + str(ventana) + r"}(" + "|".join(escaped) + r").{0," + str(ventana) + r"}", re.IGNORECASE | re.DOTALL)

def dividir_texto(texto: str, tamano: int = 2000) -> list[str]:
    return [texto[i:i+tamano] for i in range(0, len(texto), tamano)]

def detectar_idioma(texto: str, manual: str = "auto") -> str:
    if manual != "auto":
        return manual
    try:
        d = detect(texto)
        return d if d in KEYWORDS else "es"
    except:
        return "es"

def extraer_texto_archivo(archivo) -> str:
    if archivo is None:
        return ""
    nombre = archivo.name.lower()
    try:
        if nombre.endswith(".txt"):
            with open(archivo.name, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        elif nombre.endswith(".pdf"):
            if not PDF_READ:
                return "⚠️ pdfplumber no instalado. Instala con: pip install pdfplumber"
            texto = ""
            with pdfplumber.open(archivo.name) as pdf:
                for page in pdf.pages:
                    t = page.extract_text()
                    if t:
                        texto += t + "\n"
            return texto
        elif nombre.endswith(".docx"):
            if not DOCX_SUPPORT:
                return "⚠️ python-docx no instalado. Instala con: pip install python-docx"
            doc = DocxDocument(archivo.name)
            return "\n".join(p.text for p in doc.paragraphs)
    except Exception as e:
        return f"⚠️ Error al leer archivo: {e}"
    return ""

def extract_clauses(texto: str, lang: str, extra_kw: dict | None = None) -> tuple[dict, list]:
    kw = {k: list(v) for k, v in KEYWORDS.get(lang, KEYWORDS["es"]).items()}
    if extra_kw:
        for k, v in extra_kw.items():
            kw.setdefault(k, []).extend(v)
    patrones = {k: generar_regex(v) for k, v in kw.items()}
    frases = [f.strip() for f in re.split(r'\. |\.\n', texto) if f.strip()]
    clausulas = {k: [] for k in patrones}
    for i, f in enumerate(frases):
        for key, pat in patrones.items():
            if pat.search(f):
                clausulas[key].append({"ref": i + 1, "text": f})
    return clausulas, frases

def clasificar_riesgo(frase: str, lang: str) -> str:
    rmap = RIESGOS.get(lang, RIESGOS["es"])
    for nivel in ["Crítico", "Alto", "Moderado", "Bajo"]:
        for w in rmap[nivel]:
            if w.lower() in frase.lower():
                return nivel
    return ""

def detectar_riesgos(frases: list, lang: str) -> list:
    out = []
    for i, f in enumerate(frases):
        nivel = clasificar_riesgo(f, lang)
        if nivel:
            out.append({"ref": i + 1, "text": f, "nivel": f"{RISK_ICONS[nivel]} {nivel}", "nivel_raw": nivel})
    return out

def calcular_score(riesgos: list, pesos: dict | None = None) -> int:
    pesos = pesos or PESOS_DEFAULT
    return sum(pesos.get(r.get("nivel_raw", r["nivel"].split()[-1]), 0) for r in riesgos)

def label_score(score: int) -> str:
    if score <= 3:   return "🟢 Bajo"
    if score <= 6:   return "🟡 Moderado"
    if score <= 9:   return "🔴 Alto"
    return "💀 Crítico"

def extraer_fechas(texto: str) -> list[str]:
    out = []
    for p in DATE_PATTERNS:
        out += [m.group().strip() for m in re.finditer(p, texto, re.IGNORECASE)]
    return list(dict.fromkeys(out))

def extraer_montos(texto: str) -> list[str]:
    out = []
    for p in MONEY_PATTERNS:
        out += [m.group().strip() for m in re.finditer(p, texto, re.IGNORECASE)]
    return list(dict.fromkeys(out))

def extraer_partes(texto: str) -> list[str]:
    def _norm_key(s):
        s = _CLEAN_PREFIX.sub("", s).strip()
        return re.sub(r'[\s\.]+$', '', s).lower()

    found = []
    seen: set = set()
    for p in PARTY_PATTERNS:
        for m in re.findall(p, texto, re.IGNORECASE):
            parts = [m] if isinstance(m, str) else list(m)
            for raw in parts:
                raw = _CLEAN_PREFIX.sub("", raw).strip().strip('"').strip()
                raw = re.sub(r'[\s\.]+$', '', raw).strip()
                if len(raw) < 3:
                    continue
                key = _norm_key(raw)
                if any(key == s or key in s or s in key for s in seen):
                    continue
                seen.add(key)
                found.append(raw)
    return found[:8]

def detectar_abusivas(texto: str, lang: str) -> list:
    patrones = ABUSIVAS.get(lang, ABUSIVAS["es"])
    frases = [f.strip() for f in re.split(r'\. |\.\n', texto) if f.strip()]
    out = []
    for i, frase in enumerate(frases):
        for patron in patrones:
            if patron.lower() in frase.lower():
                out.append({"ref": i + 1, "text": frase, "patron": patron})
    return out

def estadisticas(texto: str) -> dict:
    palabras = len(texto.split())
    return {
        "Palabras": palabras,
        "Caracteres": len(texto),
        "Frases": len(re.split(r'\. |\.\n', texto)),
        "Páginas est.": max(1, round(palabras / 250)),
    }

def generar_checklist(clausulas: dict, plantilla: list | None = None) -> dict:
    plantilla = plantilla or list(KEYWORDS["es"].keys())
    return {c: "✅" if clausulas.get(c) else "✗" for c in plantilla}

# ── NEW: Feature 4 — Vigencia / vencimiento ───────────────────────────────────

def _parse_fecha(fecha_str: str, lang: str):
    fecha_str = fecha_str.strip().lower()
    m = re.match(r'(\d{4})-(\d{2})-(\d{2})', fecha_str)
    if m:
        try:
            return datetime(int(m.group(1)), int(m.group(2)), int(m.group(3)))
        except ValueError:
            pass
    m = re.match(r'(\d{1,2})/(\d{1,2})/(\d{2,4})', fecha_str)
    if m:
        d, mo, y = int(m.group(1)), int(m.group(2)), int(m.group(3))
        if y < 100:
            y += 2000
        try:
            return datetime(y, mo, d)
        except ValueError:
            pass
    m = re.match(r'(\d{1,2})\s+de\s+(\w+)\s+de\s+(\d{4})', fecha_str)
    if m:
        d, mes_name, y = int(m.group(1)), m.group(2), int(m.group(3))
        for lg in ("es", "fr", "pt"):
            if mes_name in MESES[lg]:
                try:
                    return datetime(y, MESES[lg][mes_name], d)
                except ValueError:
                    pass
    m = re.match(r'(\w+)\s+(\d{1,2}),?\s+(\d{4})', fecha_str)
    if m:
        mes_name, d, y = m.group(1), int(m.group(2)), int(m.group(3))
        if mes_name in MESES["en"]:
            try:
                return datetime(y, MESES["en"][mes_name], d)
            except ValueError:
                pass
    return None

def detectar_vigencia(texto: str, lang: str) -> dict | None:
    kws = VIGENCIA_KEYWORDS.get(lang, VIGENCIA_KEYWORDS["es"])
    frases = [f.strip() for f in re.split(r'\. |\.\n', texto) if f.strip()]
    for f in frases:
        fl = f.lower()
        if any(k in fl for k in kws):
            for fecha_str in extraer_fechas(f):
                dt = _parse_fecha(fecha_str, lang)
                if dt:
                    dias = (dt.date() - datetime.now().date()).days
                    return {"fecha": fecha_str, "fecha_dt": dt, "dias_restantes": dias, "frase": f}
    return None

# Definiciones ─────────────────────────────────────────────

def extraer_definiciones(texto: str) -> list[dict]:
    m = re.search(r'(?:definiciones|definitions|défini[cç][aã]?o(?:es)?)\s*[:\.]?\s*\n?(.{0,3000})',
                   texto, re.IGNORECASE | re.DOTALL)
    bloque = m.group(1) if m else texto[:3000]
    patrones = [
        r'"([^"]{2,60})"\s*(?:significa|se entiende por|quiere decir|means|shall mean|refers to|signifie)\s+([^.\n]{5,250})',
        r'«([^»]{2,60})»\s*(?:significa|se entiende por)\s+([^.\n]{5,250})',
    ]
    out, seen = [], set()
    for p in patrones:
        for term, definicion in re.findall(p, bloque, re.IGNORECASE):
            key = term.strip().lower()
            if key in seen:
                continue
            seen.add(key)
            out.append({"termino": term.strip(), "definicion": definicion.strip()})
    return out[:20]

# Contradicciones internas ─────────────────────────────────

def detectar_contradicciones(clausulas: dict) -> list[dict]:
    contradicciones = []
    for ctype, items in clausulas.items():
        valores: dict[str, list] = {}
        for it in items:
            for m in DURATION_PATTERN.finditer(it["text"]):
                val = f"{m.group(1)} {m.group(2).lower()}"
                valores.setdefault(val, []).append(it["ref"])
        if len(valores) > 1:
            contradicciones.append({"tipo": ctype, "valores": valores})
    return contradicciones

# Legibilidad (Fernández-Huerta, aproximado) ─────────────

_VOWELS_RE = re.compile(r'[aeiouáéíóúüAEIOUÁÉÍÓÚÜ]+')

def _contar_silabas(palabra: str) -> int:
    return max(1, len(_VOWELS_RE.findall(palabra)))

def calcular_legibilidad(texto: str) -> dict:
    frases = [f for f in re.split(r'[.!?\n]+', texto) if f.strip()]
    palabras = re.findall(r"[A-Za-zÁÉÍÓÚÑáéíóúñÜü']+", texto)
    n_frases = max(1, len(frases))
    n_palabras = max(1, len(palabras))
    n_silabas = sum(_contar_silabas(p) for p in palabras)
    palabras_largas = sum(1 for p in palabras if len(p) > 6)
    score = 206.84 - 0.60 * (n_silabas / n_palabras * 100) - 1.02 * (n_palabras / n_frases)
    score = round(max(0, min(100, score)), 1)
    if score >= 80:   nivel = "🟢 Muy fácil"
    elif score >= 65: nivel = "🟢 Fácil"
    elif score >= 50: nivel = "🟡 Normal"
    elif score >= 30: nivel = "🔴 Difícil"
    else:             nivel = "💀 Muy difícil"
    return {
        "score": score, "nivel": nivel,
        "palabras_por_frase": round(n_palabras / n_frases, 1),
        "pct_palabras_largas": round(palabras_largas / n_palabras * 100, 1),
    }

# ═══════════════════════════════════════════════════════════════════════════════
# HTML DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════

def _tag(items, color): return "".join(
    f"<span style='display:inline-block;margin:3px 2px;padding:3px 10px;border-radius:12px;background:{color};color:white;font-size:12px'>{html.escape(str(i))}</span>"
    for i in items) or "<span style='color:#9ca3af'>Ninguno detectado</span>"

def generar_dashboard(clausulas, riesgos, checklist, score, abusivas, fechas, montos, partes, stats, lang,
                       vigencia=None, definiciones=None, contradicciones=None, legibilidad=None):
    score_label = label_score(score)
    score_color = RISK_COLORS.get(score_label.split()[-1], "#6b7280")
    score_pct   = min(100, score * 4)

    # Stats bar
    stats_html = "".join(
        f"<div style='flex:1;text-align:center;padding:12px;border-right:1px solid #f0f0f0'>"
        f"<div style='font-size:28px;font-weight:800;color:#1e293b'>{v}</div>"
        f"<div style='font-size:11px;color:#94a3b8;text-transform:uppercase;letter-spacing:.05em'>{k}</div></div>"
        for k, v in stats.items()
    )

    # Checklist (with tooltip explaining why a missing clause matters)
    def _why_esc(k):
        return html.escape(WHY_LEGAL.get(k, ""))
    checklist_html = "".join(
        f"<div title='{_why_esc(k)}' "
        f"style='display:inline-flex;align-items:center;gap:6px;margin:4px;padding:5px 14px;"
        f"border-radius:20px;background:{'#dcfce7' if v=='✅' else '#fee2e2'};"
        f"color:{'#166534' if v=='✅' else '#991b1b'};font-size:13px;font-weight:600;cursor:help'>"
        f"{v} {CLAUSE_EMOJIS.get(k,'📄')} {k.replace('_',' ').capitalize()}</div>"
        for k, v in checklist.items()
    )
    checklist_missing_notes = "".join(
        f"<div style='font-size:12px;color:#991b1b;margin:3px 0'>✗ <b>{k.replace('_',' ').capitalize()}</b>: {html.escape(WHY_LEGAL.get(k,''))}</div>"
        for k, v in checklist.items() if v == "✗"
    )

    # Vigencia badge
    if vigencia:
        dias = vigencia["dias_restantes"]
        if dias < 0:
            vcolor, vtext = "#7c3aed", f"⏰ Contrato vencido hace {abs(dias)} días ({vigencia['fecha']})"
        elif dias <= 30:
            vcolor, vtext = "#ef4444", f"⏰ Vence en {dias} días ({vigencia['fecha']})"
        elif dias <= 90:
            vcolor, vtext = "#f59e0b", f"📅 Vence en {dias} días ({vigencia['fecha']})"
        else:
            vcolor, vtext = "#22c55e", f"📅 Vence en {dias} días ({vigencia['fecha']})"
        vigencia_html = (
            f"<div style='background:{vcolor}15;border-left:4px solid {vcolor};padding:10px 16px;"
            f"border-radius:8px;font-size:13px;color:{vcolor};font-weight:700;margin-bottom:16px'>{vtext}</div>"
        )
    else:
        vigencia_html = (
            "<div style='background:#f1f5f9;border-left:4px solid #94a3b8;padding:10px 16px;"
            "border-radius:8px;font-size:13px;color:#64748b;margin-bottom:16px'>"
            "📅 No se detectó una fecha de vigencia/vencimiento explícita.</div>"
        )

    # Legibilidad card
    if legibilidad:
        leg_html = (
            f"<div style='background:white;border-radius:12px;padding:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>"
            f"<h4 style='margin:0 0 10px;font-size:13px;color:#1e293b'>📖 Legibilidad (aprox.)</h4>"
            f"<div style='font-size:22px;font-weight:800;color:#1e293b'>{legibilidad['score']} <span style='font-size:13px;font-weight:600'>{legibilidad['nivel']}</span></div>"
            f"<div style='font-size:11px;color:#94a3b8;margin-top:6px'>{legibilidad['palabras_por_frase']} palabras/frase · "
            f"{legibilidad['pct_palabras_largas']}% palabras largas</div></div>"
        )
    else:
        leg_html = ""

    # Definiciones
    if definiciones:
        def_html = "".join(
            f"<div style='padding:6px 0;border-bottom:1px solid #f1f5f9;font-size:12px'>"
            f"<b style='color:#1e293b'>{html.escape(d['termino'])}</b>: "
            f"<span style='color:#64748b'>{html.escape(d['definicion'][:150])}</span></div>"
            for d in definiciones
        )
    else:
        def_html = "<div style='color:#9ca3af;font-size:13px'>No se detectó una sección de definiciones.</div>"

    # Contradicciones
    def _valores_str(valores):
        partes = []
        for v, refs in valores.items():
            refs_str = "/".join(str(r) for r in refs)
            partes.append(f"{v} (Ref {refs_str})")
        return ", ".join(partes)

    if contradicciones:
        contra_html = "".join(
            f"<div style='background:#fef2f2;border-left:4px solid #dc2626;padding:8px 12px;"
            f"margin:5px 0;border-radius:6px;font-size:12px;color:#7f1d1d'>"
            f"⚠️ <b>{c['tipo'].replace('_',' ').capitalize()}</b>: se detectaron valores distintos → "
            f"{_valores_str(c['valores'])}</div>"
            for c in contradicciones
        )
    else:
        contra_html = "<div style='color:#16a34a;font-size:13px'>✅ Sin contradicciones evidentes en plazos.</div>"

    # Clause cards
    cards_html = ""
    for key, items in clausulas.items():
        c = CLAUSE_COLORS.get(key, "#6b7280")
        e = CLAUSE_EMOJIS.get(key, "📄")
        rows = "".join(
            f"<div style='padding:6px 0;border-bottom:1px solid #f1f5f9;font-size:12px;color:#475569'>"
            f"<span style='background:{c}22;color:{c};border-radius:4px;padding:1px 6px;font-size:10px;font-weight:700'>Ref {it['ref']}</span> "
            f"{html.escape(it['text'][:140])}{'…' if len(it['text'])>140 else ''}</div>"
            for it in items[:6]
        )
        extra = f"<div style='font-size:11px;color:#94a3b8;margin-top:4px'>+{len(items)-6} más…</div>" if len(items) > 6 else ""
        cards_html += (
            f"<div style='background:white;border-radius:12px;padding:16px;border-top:4px solid {c};"
            f"box-shadow:0 1px 4px rgba(0,0,0,.07);margin-bottom:12px'>"
            f"<div style='display:flex;align-items:center;justify-content:space-between;margin-bottom:10px'>"
            f"<span style='font-size:15px;font-weight:700;color:#1e293b'>{e} {key.replace('_',' ').capitalize()}</span>"
            f"<span style='background:{c};color:white;border-radius:20px;padding:2px 12px;font-size:13px;font-weight:700'>{len(items)}</span></div>"
            f"{'<div style=\"color:#9ca3af;font-size:13px\">No encontrado</div>' if not items else rows + extra}</div>"
        )

    # Abusive clauses
    if abusivas:
        ab_html = "".join(
            f"<div style='background:#fff7ed;border-left:4px solid #f59e0b;padding:8px 12px;"
            f"margin:5px 0;border-radius:6px;font-size:12px;color:#78350f'>"
            f"⚠️ <b>«{html.escape(ab['patron'])}»</b> "
            f"<span style='color:#92400e'>[Ref {ab['ref']}]</span> — "
            f"{html.escape(ab['text'][:120])}…</div>"
            for ab in abusivas
        )
    else:
        ab_html = "<div style='color:#16a34a;font-size:13px'>✅ No se detectaron cláusulas potencialmente abusivas.</div>"

    # Risk list
    def _risk_row(r):
        nivel_key = r.get("nivel_raw", r["nivel"].split()[-1])
        color = RISK_COLORS.get(nivel_key, "#6b7280")
        return (
            f"<div style='background:white;border-left:4px solid {color};"
            f"padding:8px 12px;margin:5px 0;border-radius:6px;font-size:12px;"
            f"box-shadow:0 1px 3px rgba(0,0,0,.05)'>"
            f"<span style='font-weight:700'>[Ref {r['ref']}]</span> "
            f"{html.escape(r['text'][:160])} <b>→ {r['nivel']}</b></div>"
        )
    risk_html = "".join(_risk_row(r) for r in riesgos) or \
        "<div style='color:#16a34a;font-size:13px'>✅ Sin riesgos detectados.</div>"

    return f"""
<div style='font-family:"Segoe UI",system-ui,sans-serif;background:#f8fafc;padding:20px;border-radius:16px;max-width:960px;margin:0 auto'>

  <!-- Header -->
  <div style='background:linear-gradient(135deg,#0f172a 0%,#1e40af 100%);color:white;border-radius:14px;padding:24px 28px;margin-bottom:20px;display:flex;align-items:center;justify-content:space-between'>
    <div>
      <h2 style='margin:0;font-size:22px;font-weight:800;letter-spacing:-.03em'>📑 Contract Analyzer AI</h2>
      <p style='margin:4px 0 0;opacity:.65;font-size:13px'>Informe automático · {datetime.now().strftime("%d/%m/%Y %H:%M")} · Idioma: <b>{lang.upper()}</b></p>
    </div>
    <div style='background:rgba(255,255,255,.15);border-radius:10px;padding:10px 18px;text-align:center'>
      <div style='font-size:26px;font-weight:900'>{score}</div>
      <div style='font-size:11px;opacity:.8'>SCORE</div>
    </div>
  </div>

  <!-- Vigencia -->
  {vigencia_html}

  <!-- Stats -->
  <div style='background:white;border-radius:12px;display:flex;margin-bottom:16px;box-shadow:0 1px 4px rgba(0,0,0,.06);overflow:hidden'>
    {stats_html}
  </div>

  <!-- Score bar -->
  <div style='background:white;border-radius:12px;padding:18px 22px;margin-bottom:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
    <div style='display:flex;justify-content:space-between;align-items:center;margin-bottom:10px'>
      <span style='font-weight:700;color:#1e293b'>📊 Score Global de Riesgo</span>
      <span style='font-size:18px;font-weight:800;color:{score_color}'>{score} → {score_label}</span>
    </div>
    <div style='background:#e2e8f0;border-radius:20px;height:10px'>
      <div style='background:{score_color};width:{score_pct}%;height:10px;border-radius:20px'></div>
    </div>
  </div>

  <!-- Checklist -->
  <div style='background:white;border-radius:12px;padding:18px 22px;margin-bottom:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
    <h3 style='margin:0 0 12px;font-size:15px;color:#1e293b'>📋 Checklist Legal <span style='font-size:11px;color:#94a3b8;font-weight:400'>(pasa el cursor sobre cada una)</span></h3>
    {checklist_html}
    <div style='margin-top:10px'>{checklist_missing_notes}</div>
  </div>

  <!-- Partes / Fechas / Montos / Legibilidad -->
  <div style='display:grid;grid-template-columns:repeat(auto-fit,minmax(200px,1fr));gap:14px;margin-bottom:16px'>
    <div style='background:white;border-radius:12px;padding:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
      <h4 style='margin:0 0 10px;font-size:13px;color:#1e293b'>👥 Partes Identificadas</h4>
      {_tag(partes, "#8b5cf6")}
    </div>
    <div style='background:white;border-radius:12px;padding:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
      <h4 style='margin:0 0 10px;font-size:13px;color:#1e293b'>📅 Fechas y Plazos</h4>
      {_tag(fechas[:12], "#3b82f6")}
    </div>
    <div style='background:white;border-radius:12px;padding:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
      <h4 style='margin:0 0 10px;font-size:13px;color:#1e293b'>💰 Montos Detectados</h4>
      {_tag(montos[:12], "#10b981")}
    </div>
    {leg_html}
  </div>

  <!-- Definiciones -->
  <div style='background:white;border-radius:12px;padding:18px 22px;margin-bottom:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
    <h3 style='margin:0 0 12px;font-size:15px;color:#1e293b'>📚 Definiciones del Contrato</h3>
    {def_html}
  </div>

  <!-- Contradicciones -->
  <div style='background:white;border-radius:12px;padding:18px 22px;margin-bottom:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
    <h3 style='margin:0 0 12px;font-size:15px;color:#1e293b'>🔀 Posibles Contradicciones Internas</h3>
    {contra_html}
  </div>

  <!-- Abusivas -->
  <div style='background:white;border-radius:12px;padding:18px 22px;margin-bottom:16px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
    <h3 style='margin:0 0 12px;font-size:15px;color:#1e293b'>🚫 Cláusulas Potencialmente Abusivas</h3>
    {ab_html}
  </div>

  <!-- Clauses -->
  <div style='margin-bottom:16px'>
    <h3 style='font-size:15px;color:#1e293b;margin:0 0 12px'>📂 Cláusulas por Tipo</h3>
    {cards_html}
  </div>

  <!-- Risks -->
  <div style='background:white;border-radius:12px;padding:18px 22px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
    <h3 style='margin:0 0 12px;font-size:15px;color:#1e293b'>🚨 Riesgos Detectados</h3>
    {risk_html}
  </div>
</div>"""

# ═══════════════════════════════════════════════════════════════════════════════
# Vista con resaltado inline sobre el texto completo
# ═══════════════════════════════════════════════════════════════════════════════

def generar_texto_resaltado(frases, clausulas, riesgos, abusivas, lang) -> str:
    ref_to_clauses: dict[int, list[str]] = {}
    for ctype, items in clausulas.items():
        for it in items:
            ref_to_clauses.setdefault(it["ref"], []).append(ctype)
    ref_to_risk = {r["ref"]: r for r in riesgos}
    abusivas_refs = {ab["ref"] for ab in abusivas}

    partes_html = []
    for i, frase in enumerate(frases):
        ref = i + 1
        texto_escaped = html.escape(frase)
        tags = ref_to_clauses.get(ref, [])
        risk = ref_to_risk.get(ref)
        is_abusiva = ref in abusivas_refs
        style_extra = "border-bottom:2px dashed #f59e0b;" if is_abusiva else ""

        if risk:
            color = RISK_COLORS[risk["nivel_raw"]]
            tooltip = f"Riesgo: {risk['nivel']}" + (f" · También: {', '.join(t.replace('_',' ') for t in tags)}" if tags else "")
            partes_html.append(
                f"<span title='{html.escape(tooltip)}' style='background:{color}26;border-left:3px solid {color};"
                f"padding:2px 5px;border-radius:3px;{style_extra}cursor:help'>{texto_escaped}</span>. "
            )
        elif tags:
            color = CLAUSE_COLORS.get(tags[0], "#6b7280")
            tooltip = " · ".join(t.replace("_", " ").capitalize() for t in tags)
            partes_html.append(
                f"<span title='{html.escape(tooltip)}' style='background:{color}1f;border-left:3px solid {color};"
                f"padding:2px 5px;border-radius:3px;{style_extra}cursor:help'>{texto_escaped}</span>. "
            )
        else:
            partes_html.append(f"<span style='{style_extra}'>{texto_escaped}</span>. ")

    body = "".join(partes_html)

    legend_items = "".join(
        f"<span style='display:inline-flex;align-items:center;gap:5px;margin:4px 10px 4px 0;font-size:12px;color:#475569'>"
        f"<span style='width:12px;height:12px;border-radius:3px;background:{c};display:inline-block'></span>{k.replace('_',' ').capitalize()}</span>"
        for k, c in CLAUSE_COLORS.items()
    )
    legend_riesgo = "".join(
        f"<span style='display:inline-flex;align-items:center;gap:5px;margin:4px 10px 4px 0;font-size:12px;color:#475569'>"
        f"<span style='width:12px;height:12px;border-radius:3px;background:{c};display:inline-block'></span>Riesgo {k}</span>"
        for k, c in RISK_COLORS.items()
    )

    return f"""
<div style='font-family:"Segoe UI",system-ui,sans-serif;background:#f8fafc;padding:20px;border-radius:16px;max-width:960px;margin:0 auto'>
  <div style='background:white;border-radius:12px;padding:16px 20px;margin-bottom:14px;box-shadow:0 1px 4px rgba(0,0,0,.06)'>
    <h3 style='margin:0 0 10px;font-size:14px;color:#1e293b'>🗂️ Leyenda — Cláusulas</h3>
    <div>{legend_items}</div>
    <h3 style='margin:14px 0 10px;font-size:14px;color:#1e293b'>🚨 Leyenda — Riesgos <span style='font-weight:400;font-size:11px;color:#94a3b8'>(prioridad sobre cláusulas)</span></h3>
    <div>{legend_riesgo}</div>
    <div style='margin-top:10px;font-size:12px;color:#92400e'>┄ Subrayado punteado naranja = posible cláusula abusiva</div>
  </div>
  <div style='background:white;border-radius:12px;padding:22px 26px;box-shadow:0 1px 4px rgba(0,0,0,.06);line-height:1.9;font-size:14px;color:#1e293b'>
    {body}
  </div>
</div>"""

# ═══════════════════════════════════════════════════════════════════════════════
# CHARTS
# ═══════════════════════════════════════════════════════════════════════════════

def generar_graficos(clausulas: dict, riesgos: list) -> str:
    fig, (ax1, ax2) = plt.subplots(1, 2, figsize=(11, 4.5))
    fig.patch.set_facecolor("#f8fafc")

    # Bar – clauses per type
    names  = [k.replace("_", " ").capitalize() for k in clausulas]
    counts = [len(v) for v in clausulas.values()]
    colors = [CLAUSE_COLORS.get(k, "#6b7280") for k in clausulas]
    bars = ax1.bar(names, counts, color=colors, edgecolor="white", linewidth=1.5, width=0.6)
    ax1.set_facecolor("#f8fafc")
    ax1.set_title("Cláusulas por tipo", fontweight="bold", fontsize=13, pad=12)
    ax1.set_ylabel("Cantidad", fontsize=11)
    ax1.tick_params(axis="x", rotation=35, labelsize=8)
    ax1.spines[["top","right"]].set_visible(False)
    for bar, val in zip(bars, counts):
        ax1.text(bar.get_x() + bar.get_width()/2, val + .05, str(val),
                 ha="center", fontweight="bold", fontsize=11)

    # Pie – risk distribution
    niveles = {"Bajo": 0, "Moderado": 0, "Alto": 0, "Crítico": 0}
    for r in riesgos:
        k = r.get("nivel_raw", r["nivel"].split()[-1])
        niveles[k] = niveles.get(k, 0) + 1

    labels = [k for k, v in niveles.items() if v > 0]
    sizes  = [v for v in niveles.values() if v > 0]
    pie_colors = [RISK_COLORS[k] for k in labels]

    ax2.set_facecolor("#f8fafc")
    if sizes:
        wedges, texts, autotexts = ax2.pie(
            sizes, labels=labels, colors=pie_colors,
            autopct="%1.0f%%", startangle=140,
            wedgeprops={"edgecolor":"white","linewidth":2})
        for at in autotexts:
            at.set_fontweight("bold")
    else:
        ax2.text(0, 0, "Sin riesgos\ndetectados", ha="center", va="center",
                 fontsize=14, color="#16a34a", fontweight="bold")
        ax2.axis("off")
    ax2.set_title("Distribución de Riesgos", fontweight="bold", fontsize=13, pad=12)

    plt.tight_layout(pad=2)
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
    plt.savefig(tmp.name, dpi=120, bbox_inches="tight", facecolor="#f8fafc")
    plt.close()
    return tmp.name

# ═══════════════════════════════════════════════════════════════════════════════
# COMPARATOR (side-by-side, stats, diff-only toggle)
# ═══════════════════════════════════════════════════════════════════════════════

def _word_diff(l1: str, l2: str) -> tuple[str, str]:
    m = SequenceMatcher(None, l1.split(), l2.split())
    left = right = ""
    for tag, i1, i2, j1, j2 in m.get_opcodes():
        w1 = " ".join(l1.split()[i1:i2])
        w2 = " ".join(l2.split()[j1:j2])
        if tag == "equal":
            left  += w1 + " "
            right += w2 + " "
        elif tag == "replace":
            left  += f"<mark style='background:#fca5a5;border-radius:3px'>{html.escape(w1)}</mark> "
            right += f"<mark style='background:#86efac;border-radius:3px'>{html.escape(w2)}</mark> "
        elif tag == "delete":
            left  += f"<mark style='background:#fca5a5;border-radius:3px;text-decoration:line-through'>{html.escape(w1)}</mark> "
        elif tag == "insert":
            right += f"<mark style='background:#86efac;border-radius:3px'>{html.escape(w2)}</mark> "
    return left.strip(), right.strip()

def comparar_contratos(a: str, b: str, solo_diffs: bool = False) -> str:
    if not a or not b:
        return "<p style='color:#ef4444;font-family:sans-serif'>⚠️ Introduce ambos contratos.</p>"

    lines1 = a.splitlines()
    lines2 = b.splitlines()
    s = SequenceMatcher(None, lines1, lines2)
    ratio   = round(s.ratio() * 100, 1)
    added = deleted = modified = 0

    col_a = col_b = ""
    ROW_BASE = "padding:3px 8px;font-size:12px;font-family:monospace;white-space:pre-wrap;word-break:break-word;"

    for tag, i1, i2, j1, j2 in s.get_opcodes():
        if tag == "equal":
            if not solo_diffs:
                for l in lines1[i1:i2]:
                    safe = html.escape(l) or "&nbsp;"
                    col_a += f"<div style='{ROW_BASE}'>{safe}</div>"
                    col_b += f"<div style='{ROW_BASE}'>{safe}</div>"
        elif tag == "delete":
            deleted += i2 - i1
            for l in lines1[i1:i2]:
                col_a += f"<div style='{ROW_BASE}background:#fee2e2'>{html.escape(l) or '&nbsp;'}</div>"
                col_b += f"<div style='{ROW_BASE}background:#f8fafc'>&nbsp;</div>"
        elif tag == "insert":
            added += j2 - j1
            for l in lines2[j1:j2]:
                col_a += f"<div style='{ROW_BASE}background:#f8fafc'>&nbsp;</div>"
                col_b += f"<div style='{ROW_BASE}background:#dcfce7'>{html.escape(l) or '&nbsp;'}</div>"
        elif tag == "replace":
            modified += max(i2-i1, j2-j1)
            for k in range(max(i2-i1, j2-j1)):
                l1 = lines1[i1+k] if i1+k < i2 else ""
                l2 = lines2[j1+k] if j1+k < j2 else ""
                dl, dr = _word_diff(l1, l2)
                col_a += f"<div style='{ROW_BASE}background:#fef3c7'>{dl or '&nbsp;'}</div>"
                col_b += f"<div style='{ROW_BASE}background:#ecfdf5'>{dr or '&nbsp;'}</div>"

    ratio_color = "#16a34a" if ratio > 80 else "#d97706" if ratio > 50 else "#dc2626"

    badge = lambda bg, text: (
        f"<span style='background:{bg};color:white;padding:5px 14px;border-radius:20px;"
        f"font-size:13px;font-weight:700'>{text}</span>"
    )

    return f"""
<div style='font-family:"Segoe UI",system-ui,sans-serif;padding:4px'>
  <div style='display:flex;gap:10px;margin-bottom:14px;flex-wrap:wrap'>
    {badge(ratio_color, f"📊 Similitud: {ratio}%")}
    {badge("#16a34a", f"➕ Añadidas: {added}")}
    {badge("#dc2626", f"➖ Eliminadas: {deleted}")}
    {badge("#d97706", f"✏️ Modificadas: {modified}")}
  </div>
  <div style='display:grid;grid-template-columns:1fr 1fr;border:1px solid #e2e8f0;border-radius:10px;overflow:hidden'>
    <div style='background:#fef2f2;padding:8px 14px;font-weight:700;font-size:13px;color:#991b1b;border-bottom:1px solid #e2e8f0'>📄 Contrato A</div>
    <div style='background:#f0fdf4;padding:8px 14px;font-weight:700;font-size:13px;color:#166534;border-bottom:1px solid #e2e8f0;border-left:1px solid #e2e8f0'>📄 Contrato B</div>
    <div style='padding:6px;max-height:520px;overflow-y:auto'>{col_a or "<div style='padding:8px;color:#9ca3af'>Sin diferencias</div>"}</div>
    <div style='padding:6px;max-height:520px;overflow-y:auto;border-left:1px solid #e2e8f0'>{col_b or "<div style='padding:8px;color:#9ca3af'>Sin diferencias</div>"}</div>
  </div>
</div>"""

# ═══════════════════════════════════════════════════════════════════════════════
# EXPORT FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════════════

def exportar_html(md: str) -> str | None:
    if not md:
        return None
    try:
        import markdown
        content = markdown.markdown(md)
    except ImportError:
        content = md.replace("\n", "<br>")
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".html", mode="w", encoding="utf-8")
    tmp.write(f"<!DOCTYPE html><html><head><meta charset='utf-8'>"
              f"<style>body{{font-family:sans-serif;max-width:800px;margin:40px auto;padding:20px}}"
              f"h2{{color:#1e40af}}h3{{color:#1e293b}}</style></head><body>{content}</body></html>")
    tmp.close()
    return tmp.name


def exportar_csv(clausulas: dict, riesgos: list) -> str | None:
    if not clausulas:
        return None
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["Tipo", "Ref", "Texto", "Nivel de Riesgo"])
    for tipo, items in clausulas.items():
        for it in items:
            w.writerow([tipo, it["ref"], it["text"], ""])
    for r in riesgos:
        w.writerow(["RIESGO", r["ref"], r["text"], r["nivel"]])
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".csv", mode="w", encoding="utf-8", newline="")
    tmp.write(buf.getvalue()); tmp.close()
    return tmp.name


# Exportación real a Word con estilos y colores ───────────

def _set_cell_background(cell, color_hex: str):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)

def exportar_docx(resultado: dict) -> str | None:
    if not DOCX_SUPPORT or not resultado:
        return None
    doc = DocxDocument()

    title = doc.add_heading("📑 Informe de Análisis de Contrato", level=0)
    doc.add_paragraph(f"Generado el {datetime.now().strftime('%d/%m/%Y %H:%M')} · Idioma: {resultado['lang'].upper()}")

    doc.add_heading("Resumen Ejecutivo", level=1)
    p = doc.add_paragraph()
    r = p.add_run(f"Score de riesgo global: {resultado['score']} → {label_score(resultado['score'])}")
    r.bold = True
    doc.add_paragraph(f"Palabras: {resultado['stats']['Palabras']}  ·  Páginas est.: {resultado['stats']['Páginas est.']}")
    if resultado.get("vigencia"):
        v = resultado["vigencia"]
        doc.add_paragraph(f"Vigencia detectada: {v['fecha']} ({v['dias_restantes']} días restantes)")

    doc.add_heading("Checklist Legal", level=1)
    for k, v in resultado["checklist"].items():
        doc.add_paragraph(f"{v}  {k.replace('_',' ').capitalize()}", style=None)

    doc.add_heading("Cláusulas por Tipo", level=1)
    for key, items in resultado["clausulas"].items():
        h = doc.add_heading(f"{CLAUSE_EMOJIS.get(key,'📄')} {key.replace('_',' ').capitalize()}", level=2)
        for run in h.runs:
            run.font.color.rgb = RGBColor.from_string(CLAUSE_COLORS.get(key, "#333333").lstrip("#"))
        if items:
            for it in items[:15]:
                doc.add_paragraph(f"[Ref {it['ref']}] {it['text']}", style="List Bullet")
        else:
            doc.add_paragraph("No encontrado.")

    doc.add_heading("Riesgos Detectados", level=1)
    if resultado["riesgos"]:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Light Grid Accent 1"
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text = "Ref", "Texto", "Nivel"
        for r_ in resultado["riesgos"]:
            row = table.add_row().cells
            row[0].text = str(r_["ref"])
            row[1].text = r_["text"][:200]
            row[2].text = r_["nivel"]
            _set_cell_background(row[2], RISK_COLORS.get(r_["nivel_raw"], "#cccccc"))
    else:
        doc.add_paragraph("Sin riesgos detectados.")

    if resultado.get("abusivas"):
        doc.add_heading("Cláusulas Potencialmente Abusivas", level=1)
        for ab in resultado["abusivas"]:
            doc.add_paragraph(f"«{ab['patron']}» [Ref {ab['ref']}] — {ab['text'][:200]}", style="List Bullet")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


#  Exportación real a Excel con formato condicional ────────

def exportar_xlsx(resultado: dict) -> str | None:
    if not XLSX_SUPPORT or not resultado:
        return None
    wb = Workbook()

    ws1 = wb.active
    ws1.title = "Riesgos"
    ws1.append(["Ref", "Texto", "Nivel"])
    for cell in ws1[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1E40AF")
    for r_ in resultado["riesgos"]:
        ws1.append([r_["ref"], r_["text"], r_["nivel_raw"]])
        fill_color = RISK_COLORS.get(r_["nivel_raw"], "#cccccc").lstrip("#")
        ws1.cell(row=ws1.max_row, column=3).fill = PatternFill("solid", fgColor=fill_color)
    ws1.column_dimensions["B"].width = 80
    ws1.column_dimensions["A"].width = 8
    ws1.column_dimensions["C"].width = 14

    ws2 = wb.create_sheet("Cláusulas")
    ws2.append(["Tipo", "Ref", "Texto"])
    for cell in ws2[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1E40AF")
    for tipo, items in resultado["clausulas"].items():
        for it in items:
            ws2.append([tipo, it["ref"], it["text"]])
    ws2.column_dimensions["C"].width = 90
    ws2.column_dimensions["A"].width = 22

    ws3 = wb.create_sheet("Resumen")
    ws3.append(["Métrica", "Valor"])
    for cell in ws3[1]:
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = PatternFill("solid", fgColor="1E40AF")
    ws3.append(["Score global", resultado["score"]])
    ws3.append(["Nivel de riesgo", label_score(resultado["score"])])
    ws3.append(["Idioma", resultado["lang"].upper()])
    ws3.append(["Palabras", resultado["stats"]["Palabras"]])
    if resultado.get("vigencia"):
        ws3.append(["Vigencia", resultado["vigencia"]["fecha"]])
        ws3.append(["Días restantes", resultado["vigencia"]["dias_restantes"]])
    ws3.column_dimensions["A"].width = 22
    ws3.column_dimensions["B"].width = 30

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx")
    wb.save(tmp.name)
    return tmp.name

# ═══════════════════════════════════════════════════════════════════════════════
# MAIN ANALYSIS PIPELINE
# ═══════════════════════════════════════════════════════════════════════════════

def _analizar_core(texto: str, lang_manual: str, tipo_contrato: str,
                    pesos: dict, progress=None) -> dict:
    """Ejecuta el pipeline de análisis completo y devuelve un dict con todo."""
    if progress: progress(0.10, desc="Detectando idioma…")
    lang = detectar_idioma(texto, lang_manual)

    if progress: progress(0.20, desc="Extrayendo cláusulas…")
    kw_keys = list(KEYWORDS.get(lang, KEYWORDS["es"]).keys())
    clausulas_totales: dict[str, list] = {k: [] for k in kw_keys}
    frases_totales: list[str] = []
    for bloque in dividir_texto(texto):
        cl, fr = extract_clauses(bloque, lang)
        for k in clausulas_totales:
            clausulas_totales[k].extend(cl.get(k, []))
        frases_totales.extend(fr)

    if progress: progress(0.40, desc="Analizando riesgos…")
    riesgos   = detectar_riesgos(frases_totales, lang)
    abusivas  = detectar_abusivas(texto, lang)
    fechas    = extraer_fechas(texto)
    montos    = extraer_montos(texto)
    partes    = extraer_partes(texto)
    stats     = estadisticas(texto)

    plantilla = TEMPLATES.get(tipo_contrato, TEMPLATES["Genérico"])
    checklist = generar_checklist(clausulas_totales, plantilla)
    score     = calcular_score(riesgos, pesos)

    if progress: progress(0.60, desc="Detectando vigencia y contradicciones…")
    vigencia        = detectar_vigencia(texto, lang)
    contradicciones = detectar_contradicciones(clausulas_totales)
    definiciones    = extraer_definiciones(texto)
    legibilidad     = calcular_legibilidad(texto)

    return {
        "lang": lang, "clausulas": clausulas_totales, "frases": frases_totales,
        "riesgos": riesgos, "abusivas": abusivas, "fechas": fechas, "montos": montos,
        "partes": partes, "stats": stats, "checklist": checklist, "score": score,
        "vigencia": vigencia, "contradicciones": contradicciones,
        "definiciones": definiciones, "legibilidad": legibilidad,
        "tipo_contrato": tipo_contrato,
    }


def analizar_contrato(texto, archivo, lang_manual, tipo_contrato,
                       peso_bajo, peso_moderado, peso_alto, peso_critico,
                       progress=gr.Progress()):
    if archivo is not None:
        from_file = extraer_texto_archivo(archivo)
        if from_file and not from_file.startswith("⚠️"):
            texto = from_file
    if not texto or len(texto.strip()) < 40:
        vacio = "⚠️ El texto es demasiado corto o está vacío."
        return (vacio, None, None, None, None, None, None)

    pesos = {"Bajo": peso_bajo, "Moderado": peso_moderado, "Alto": peso_alto, "Crítico": peso_critico}
    resultado = _analizar_core(texto, lang_manual, tipo_contrato, pesos, progress)

    progress(0.80, desc="Generando visualizaciones…")
    dashboard = generar_dashboard(
        resultado["clausulas"], resultado["riesgos"], resultado["checklist"], resultado["score"],
        resultado["abusivas"], resultado["fechas"], resultado["montos"], resultado["partes"],
        resultado["stats"], resultado["lang"], vigencia=resultado["vigencia"],
        definiciones=resultado["definiciones"], contradicciones=resultado["contradicciones"],
        legibilidad=resultado["legibilidad"],
    )
    grafico   = generar_graficos(resultado["clausulas"], resultado["riesgos"])
    resaltado = generar_texto_resaltado(resultado["frases"], resultado["clausulas"],
                                         resultado["riesgos"], resultado["abusivas"], resultado["lang"])

    # Markdown report
    lang = resultado["lang"]
    checklist = resultado["checklist"]
    clausulas_totales = resultado["clausulas"]
    riesgos = resultado["riesgos"]
    abusivas = resultado["abusivas"]
    fechas = resultado["fechas"]
    montos = resultado["montos"]
    partes = resultado["partes"]
    stats = resultado["stats"]
    score = resultado["score"]

    clausulas_encontradas = [k for k, v in clausulas_totales.items() if v]
    partes_str = " · ".join(partes[:4]) if partes else "No identificadas"
    resumen_ai = (
        f"**Tipo de contrato:** {tipo_contrato} · **Idioma:** {lang.upper()} · "
        f"**Palabras:** {stats['Palabras']} · **Páginas est.:** {stats['Páginas est.']}\n\n"
        f"**Partes:** {partes_str}\n\n"
        f"**Cláusulas encontradas:** {', '.join(clausulas_encontradas) if clausulas_encontradas else 'Ninguna'}\n\n"
        f"**Riesgos detectados:** {len(riesgos)} "
        f"({'ninguno' if not riesgos else ', '.join(sorted(set(r['nivel'] for r in riesgos)))})\n\n"
        f"**Score global:** {score} → {label_score(score)}\n\n"
        f"**Legibilidad:** {resultado['legibilidad']['score']} → {resultado['legibilidad']['nivel']}"
    )

    md  = f"## 📑 Informe de Análisis · `{lang.upper()}` · {tipo_contrato}\n\n"
    md += f"### 📝 Resumen Ejecutivo\n{resumen_ai}\n\n"
    if resultado["vigencia"]:
        v = resultado["vigencia"]
        md += f"### ⏰ Vigencia\nFecha detectada: {v['fecha']} · Días restantes: {v['dias_restantes']}\n\n"
    md += "### 📋 Checklist Legal\n" + "\n".join(f"- {k.replace('_',' ').capitalize()}: {v}" for k, v in checklist.items()) + "\n\n"
    if resultado["definiciones"]:
        md += "### 📚 Definiciones\n" + "\n".join(f"- **{d['termino']}**: {d['definicion']}" for d in resultado["definiciones"]) + "\n\n"
    if resultado["contradicciones"]:
        md += "### 🔀 Posibles Contradicciones\n"
        for c in resultado["contradicciones"]:
            valores_str = ", ".join(f"{v} (Ref {'/'.join(str(x) for x in refs)})" for v, refs in c["valores"].items())
            md += f"- **{c['tipo'].replace('_',' ').capitalize()}**: {valores_str}\n"
        md += "\n"
    if fechas:
        md += "### 📅 Fechas y Plazos\n" + "\n".join(f"- {f}" for f in fechas) + "\n\n"
    if montos:
        md += "### 💰 Montos\n" + "\n".join(f"- {m}" for m in montos) + "\n\n"
    if partes:
        md += "### 👥 Partes\n" + "\n".join(f"- {p}" for p in partes) + "\n\n"
    for key in clausulas_totales:
        e = CLAUSE_EMOJIS.get(key, "📄")
        items = clausulas_totales[key]
        md += f"### {e} {key.replace('_',' ').capitalize()}\n"
        md += "\n".join(f"- [Ref {it['ref']}] {it['text']}" for it in items) if items else "- No encontrado"
        md += "\n\n"
    if abusivas:
        md += f"### 🚫 Cláusulas Abusivas ({len(abusivas)})\n"
        md += "\n".join(f"- [Ref {ab['ref']}] «{ab['patron']}» → {ab['text'][:150]}" for ab in abusivas) + "\n\n"
    md += "### 🚨 Riesgos\n"
    md += "\n".join(f"- [Ref {r['ref']}] {r['text']} → {r['nivel']}" for r in riesgos) or "- Sin riesgos detectados"
    md += f"\n\n### 📊 Score Global: {score} → {label_score(score)}\n"

    progress(1.0, desc="¡Listo!")
    return md, dashboard, grafico, resaltado, clausulas_totales, riesgos, resultado


# Análisis multi-archivo / cartera ─────────────────────────

def analizar_cartera(archivos, lang_manual, tipo_contrato,
                      peso_bajo, peso_moderado, peso_alto, peso_critico,
                      progress=gr.Progress()):
    if not archivos:
        return [["⚠️ No se subieron archivos", "", "", "", "", ""]]

    pesos = {"Bajo": peso_bajo, "Moderado": peso_moderado, "Alto": peso_alto, "Crítico": peso_critico}
    filas = []
    total = len(archivos)
    for idx, archivo in enumerate(archivos):
        progress((idx) / total, desc=f"Analizando {archivo.name.split('/')[-1]}…")
        texto = extraer_texto_archivo(archivo)
        nombre = archivo.name.split("/")[-1]
        if not texto or texto.startswith("⚠️") or len(texto.strip()) < 40:
            filas.append([nombre, "—", "—", 0, "⚠️ No se pudo leer", 0])
            continue
        resultado = _analizar_core(texto, lang_manual, tipo_contrato, pesos)
        n_abusivas = len(resultado["abusivas"])
        filas.append([
            nombre,
            resultado["lang"].upper(),
            resultado["stats"]["Palabras"],
            resultado["score"],
            label_score(resultado["score"]),
            n_abusivas,
        ])

    filas.sort(key=lambda r: r[3] if isinstance(r[3], int) else -1, reverse=True)
    progress(1.0, desc="¡Listo!")
    return filas

# ═══════════════════════════════════════════════════════════════════════════════
# GRADIO UI
# ═══════════════════════════════════════════════════════════════════════════════

CSS = """
.gr-button-primary { background: #1e40af !important; }
footer { display: none !important; }
"""

with gr.Blocks(theme=gr.themes.Soft(primary_hue="blue"), css=CSS) as demo:
    gr.Markdown(
        "# 🤖 Contract Analyzer \n"
        "Análisis legal automatizado"
    )

    # Shared state
    state_clausulas = gr.State(value=None)
    state_riesgos   = gr.State(value=None)
    state_resultado = gr.State(value=None)

    # ── Tab 1: Analyze ─────────────────────────────────────────────────────────
    with gr.Tab("📄 Analizar Contrato"):
        with gr.Row(equal_height=False):

            # Left panel – inputs
            with gr.Column(scale=1, min_width=300):
                archivo_input = gr.File(
                    label="📂 Subir contrato (.txt · .docx· .pdf)",
                    file_types=[".txt", ".pdf", ".docx"]
                )
                texto_input = gr.Textbox(
                    label="O pega el texto directamente",
                    lines=14,
                    placeholder="Pega aquí el texto del contrato…"
                )
                with gr.Row():
                    lang_input = gr.Dropdown(
                        choices=list(LANG_LABELS.keys()),
                        value="auto",
                        label="🌐 Idioma",
                        info="'auto' detecta automáticamente"
                    )
                    tipo_input = gr.Dropdown(
                        choices=list(TEMPLATES.keys()),
                        value="Genérico",
                        label="📑 Tipo de contrato",
                        info="Ajusta el checklist legal esperado"
                    )
                with gr.Accordion("⚙️ Ajustar pesos de riesgo (score)", open=False):
                    gr.Markdown("Sube o baja el peso de cada nivel según tu tolerancia al riesgo.")
                    with gr.Row():
                        peso_bajo_input      = gr.Slider(0, 10, value=1, step=1, label="🟢 Bajo")
                        peso_moderado_input  = gr.Slider(0, 10, value=2, step=1, label="🟡 Moderado")
                    with gr.Row():
                        peso_alto_input      = gr.Slider(0, 10, value=3, step=1, label="🔴 Alto")
                        peso_critico_input   = gr.Slider(0, 10, value=4, step=1, label="💀 Crítico")
                boton_analizar = gr.Button("🔍 Analizar", variant="primary", size="lg")
                gr.Markdown("---")
                with gr.Row():
                    boton_html = gr.Button("📄 HTML")
                    boton_csv  = gr.Button("📊 CSV")
                with gr.Row():
                    boton_docx = gr.Button("📝 Word")
                    boton_xlsx = gr.Button("📈 Excel")
                file_html = gr.File(label="Descarga HTML", visible=True)
                file_csv  = gr.File(label="Descarga CSV",  visible=True)
                file_docx = gr.File(label="Descarga Word", visible=True)
                file_xlsx = gr.File(label="Descarga Excel", visible=True)

            # Right panel – outputs
            with gr.Column(scale=2):
                with gr.Tabs():
                    with gr.Tab("📊 Dashboard"):
                        out_dashboard = gr.HTML()
                    with gr.Tab("📖 Vista con Resaltado"):
                        out_highlight = gr.HTML()
                    with gr.Tab("📝 Informe"):
                        out_markdown  = gr.Markdown()
                    with gr.Tab("📈 Gráficos"):
                        out_grafico   = gr.Image(label="Distribución")

        # Events
        boton_analizar.click(
            fn=analizar_contrato,
            inputs=[texto_input, archivo_input, lang_input, tipo_input,
                    peso_bajo_input, peso_moderado_input, peso_alto_input, peso_critico_input],
            outputs=[out_markdown, out_dashboard, out_grafico, out_highlight,
                     state_clausulas, state_riesgos, state_resultado],
        )
        boton_html.click(fn=exportar_html, inputs=out_markdown,  outputs=file_html)
        boton_csv.click(
            fn=lambda cl, ri: exportar_csv(cl, ri) if cl else None,
            inputs=[state_clausulas, state_riesgos],
            outputs=file_csv
        )
        boton_docx.click(
            fn=lambda res: exportar_docx(res) if res else None,
            inputs=[state_resultado],
            outputs=file_docx
        )
        boton_xlsx.click(
            fn=lambda res: exportar_xlsx(res) if res else None,
            inputs=[state_resultado],
            outputs=file_xlsx
        )

    # ── Tab 2: Compare ─────────────────────────────────────────────────────────
    with gr.Tab("🔍 Comparar Contratos"):
        with gr.Row():
            cont_a = gr.Textbox(label="📄 Contrato A", lines=18, placeholder="Pega el contrato A…")
            cont_b = gr.Textbox(label="📄 Contrato B", lines=18, placeholder="Pega el contrato B…")
        with gr.Row():
            solo_diffs   = gr.Checkbox(label="Mostrar solo diferencias", value=False)
            boton_comparar = gr.Button("▶ Comparar", variant="primary")
        out_diff = gr.HTML()
        boton_comparar.click(
            fn=comparar_contratos,
            inputs=[cont_a, cont_b, solo_diffs],
            outputs=out_diff
        )

    # ── Tab 3: Cartera (multi-file) ───────────────────────────────────────────
    with gr.Tab("📁 Análisis Múltiple (Cartera)"):
        gr.Markdown(
            "Sube varios contratos a la vez para obtener una tabla resumen "
            "ordenada por score de riesgo, útil para priorizar cuáles revisar primero."
        )
        with gr.Row():
            archivos_multi = gr.File(
                label="📂 Subir varios contratos (.txt · .docx · .pdf)",
                file_types=[".txt", ".pdf", ".docx"],
                file_count="multiple"
            )
        with gr.Row():
            lang_multi = gr.Dropdown(choices=list(LANG_LABELS.keys()), value="auto", label="🌐 Idioma")
            tipo_multi = gr.Dropdown(choices=list(TEMPLATES.keys()), value="Genérico", label="📑 Tipo de contrato")
        boton_cartera = gr.Button("🔍 Analizar Cartera", variant="primary")
        out_cartera = gr.Dataframe(
            headers=["Archivo", "Idioma", "Palabras", "Score", "Nivel de Riesgo", "Cláusulas Abusivas"],
            label="Resumen de la cartera (ordenado por score descendente)",
        )
        boton_cartera.click(
            fn=analizar_cartera,
            inputs=[archivos_multi, lang_multi, tipo_multi,
                    peso_bajo_input, peso_moderado_input, peso_alto_input, peso_critico_input],
            outputs=out_cartera
        )

if __name__ == "__main__":
    demo.launch()
